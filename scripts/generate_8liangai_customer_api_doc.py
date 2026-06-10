from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_MD = ROOT / "docs" / "8liangai_customer_api_doc.md"
OUTPUT_DOCX = ROOT / "docs" / "8liangai-customer-api-doc-v3.docx"
SITE_URL = "https://8liangai.com"
BUILD_DATE = "2026-05-25"


def set_run_font(run, latin: str = "Arial", east_asia: str = "Microsoft YaHei") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.15) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "0F4C81", 18, 10),
        ("Heading 2", 13, "0F4C81", 12, 7),
        ("Heading 3", 11.5, "1E3A5F", 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.15


def add_cover(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(title, after=4, line=1.0)
    run = title.add_run("8liangAI 新版接口文档")
    set_run_font(run)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("111111")

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=12, line=1.15)
    run = subtitle.add_run("客户版：覆盖图片生成、视频生成、状态查询、视频查看、真实资产、虚拟资产接口")
    set_run_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor.from_string("444444")

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    meta.style = "Table Grid"
    rows = [
        ("服务地址", SITE_URL),
        ("适用模型", "seedream4.5 / seedream5.0lite / seedance2 / sd2.0fast"),
        ("文档定位", "面向客户接入，强调怎么调用、怎么传参、怎么排查"),
        ("生成日期", BUILD_DATE),
    ]
    for idx, (k, v) in enumerate(rows):
        left = meta.cell(idx, 0)
        right = meta.cell(idx, 1)
        left.text = k
        right.text = v
        set_cell_shading(left, "EAF1F8")
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=4, line=1.1)
                for r in p.runs:
                    set_run_font(r)
                    r.font.size = Pt(10)
        for r in left.paragraphs[0].runs:
            r.font.bold = True

    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=10, after=0, line=1.15)
    r1 = note.add_run("说明：")
    set_run_font(r1)
    r1.font.bold = True
    r2 = note.add_run(" 本文已按客户调用习惯整理接口顺序，并对同一站点下不同返回结构做了区分。")
    set_run_font(r2)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def preprocess_markdown(text: str) -> str:
    return text.replace("https://8liangai.com", SITE_URL)


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    return stripped.startswith("|") and set(stripped.replace("|", "").replace("-", "").replace(":", "").replace(" ", "")) == set()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int] | None:
    if start + 1 >= len(lines):
        return None
    if not lines[start].strip().startswith("|") or not is_table_separator(lines[start + 1]):
        return None
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        row = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(row)
        idx += 1
    if len(rows) < 2:
        return None
    return [rows[0], *rows[2:]], idx


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            value = row[c_idx] if c_idx < len(row) else ""
            cell = table.cell(r_idx, c_idx)
            cell.text = value
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, "EAF1F8")
            for p in cell.paragraphs:
                set_paragraph_spacing(p, after=3, line=1.1)
                for run in p.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.5)
                    if r_idx == 0:
                        run.font.bold = True
    doc.add_paragraph()


def add_code_block(doc: Document, code: str, lang: str) -> None:
    if lang:
        label = doc.add_paragraph()
        set_paragraph_spacing(label, before=4, after=2, line=1.0)
        run = label.add_run(lang.upper())
        set_run_font(run)
        run.font.size = Pt(8.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string("666666")

    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=0, after=0, line=1.0)
        p.paragraph_format.left_indent = Inches(0.18)
        run = p.add_run(line)
        set_run_font(run, latin="Consolas", east_asia="Microsoft YaHei")
        run.font.size = Pt(9)
        set_paragraph_shading(p, "F5F7FA")
    doc.add_paragraph()


def add_paragraph_text(doc: Document, text: str, style: str | None = None) -> None:
    p = doc.add_paragraph(style=style)
    set_paragraph_spacing(p, before=0, after=6, line=1.2)
    run = p.add_run(text)
    set_run_font(run)


def build_doc(md_text: str) -> Document:
    doc = Document()
    style_document(doc)
    add_cover(doc)

    lines = preprocess_markdown(md_text).splitlines()
    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, "\n".join(code_lines), code_lang)
                in_code = False
                code_lang = ""
                code_lines = []
            else:
                code_lines.append(line.rstrip("\n"))
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            code_lang = stripped[3:].strip()
            i += 1
            continue

        parsed_table = parse_table(lines, i)
        if parsed_table is not None:
            rows, i = parsed_table
            add_markdown_table(doc, rows)
            continue

        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            add_paragraph_text(doc, text, style=f"Heading {level}")
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            add_paragraph_text(doc, re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            i += 1
            continue

        if stripped.startswith("- "):
            add_paragraph_text(doc, stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        add_paragraph_text(doc, stripped)
        i += 1

    return doc


def main() -> None:
    md_text = SOURCE_MD.read_text(encoding="utf-8")
    doc = build_doc(md_text)
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
